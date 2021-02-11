'''
-----------------------------------------------------------------------
Created by Manasvi Mohan Sharma on 11/01/21 (dd/mm/yy)
Project Name: ML_Equity_Signals | File Name: custom_functions.py
IDE: PyCharm | Python Version: 3.8
-----------------------------------------------------------------------
                                       _ 
                                      (_)
 _ __ ___   __ _ _ __   __ _ _____   ___ 
| '_ ` _ \ / _` | '_ \ / _` / __\ \ / / |
| | | | | | (_| | | | | (_| \__ \\ V /| |
|_| |_| |_|\__,_|_| |_|\__,_|___/ \_/ |_|

GitHub:   https://github.com/manasvimohan
Linkedin: https://www.linkedin.com/in/manasvi-mohan-sharma-119375168/
Website:  https://www.manasvi.co.in/
-----------------------------------------------------------------------
Project Information:
This uses classification approach to create buy/sell/hold signal
for any investment instrument with time series as input data, and
uses SKlearn library to model the series and predict there after
based on saved model

About this file:
This has all the custom functions required to run the main file.
-----------------------------------------------------------------------
'''

'''
Useful links
https://machinelearningmastery.com/smote-oversampling-for-imbalanced-classification/
https://towardsdatascience.com/methods-for-dealing-with-imbalanced-data-5b761be45a18/
https://www.analyticsvidhya.com/blog/2020/07/10-techniques-to-deal-with-class-imbalance-in-machine-learning/
'''

import numpy as np
import pandas as pd
from pandas_datareader import data as web

# LINK: https://pypi.org/project/imbalanced-learn/
from imblearn.under_sampling import RandomUnderSampler, TomekLinks
from imblearn.over_sampling import RandomOverSampler, SMOTE
from imblearn.pipeline import Pipeline

import pickle
from datetime import datetime

from docx import Document  # Making and exporting docx
from docx.shared import Mm  # Page setup of docx
from docx.shared import Pt
# import mammoth  # Converting docx to html

def make_time_x(data, use_of_price, lookbacklen):

    #Create a new dataframe with one feature column
    df = data.filter([use_of_price])
    #Convert the dataframe to a numpy array
    dataset = df.values
    #Split the data into x_train and y_train data sets
    x = []

    for i in range(lookbacklen, len(dataset)):
        x.append(dataset[i-lookbacklen:i, 0])

    #Convert the x_train and y_train to numpy arrays
    x = np.array(x)

    dfx = pd.DataFrame(x)
    dfx[use_of_price] = dfx[lookbacklen-1].shift(-1)
    minus_len = -1*dfx.shape[0]
    dfx.set_index(data.index[minus_len:], inplace=True, drop=True)
    dfx.dropna(inplace = True)
    return dfx


def tag_data_variable_level(data, calc_variable_name, movement):
    conditions = [(data[calc_variable_name] > movement),
                  (data[calc_variable_name] < -1*movement)]
    choices = ['Buy','Sell']
    choice_name = "choice_"+ calc_variable_name
    final = data.copy()
    final[choice_name] = np.select(conditions, choices, default='Hold')
    return final, choice_name

def make_df_OHLC_symbol(symbolname,start_date,end_date):

    df = web.DataReader(symbolname,data_source = 'yahoo', start = start_date, end = end_date)

    df.rename(columns = {'Close': 'close',
                        'Open': 'open',
                        'High': 'high',
                        'Low': 'low',
                        'Volume':'volume'}, inplace =True)
    return df


def sample_data(x, y, choice):
    seed = 42
    k = 8

    if choice == 'both':
        # over = SMOTE(sampling_strategy='auto', k_neighbors=k, random_state=seed)
        # under = RandomUnderSampler(sampling_strategy='auto')

        # strategy = {0: 1000, 1: 1000, 2: 1000}
        # over = RandomOverSampler(sampling_strategy=strategy, random_state=seed)
        # under = TomekLinks(sampling_strategy='majority')
        choice = 'Over_and_Under_Sampling'

        over = RandomOverSampler(sampling_strategy=0.5, random_state=seed)
        under = TomekLinks(sampling_strategy=0.5)

        steps = [('o', over),('u', under)]
        pipeline = Pipeline(steps=steps)
        x, y = pipeline.fit_resample(x, y)
    elif choice == 'u':
        choice = 'Under_Sampling'
        print('Performing Random Under Sample')
        strategy = 'auto'
        under = RandomUnderSampler(sampling_strategy=strategy)
        steps = [('u', under)]
        pipeline = Pipeline(steps=steps)
        x, y = pipeline.fit_resample(x, y)
    elif choice == 'o':
        choice = 'Over_Sampling'
        print('Performing Random Over Sample')
        over = RandomOverSampler(random_state=seed)
        steps = [('o', over)]
        pipeline = Pipeline(steps=steps)
        x, y = pipeline.fit_resample(x, y)
    elif choice == 'smote':
        choice = 'SMOTE'
        strategy = 'auto'
        smote_over_sample = SMOTE(sampling_strategy=strategy, k_neighbors=k, random_state=seed)
        x, y = smote_over_sample.fit_resample(x, y)

    return x, y, choice

def count_y(y,y_index):
    total_decisions = len(y[0])
    if total_decisions==3:
        if y_index==0:
            y_is = "Buy"
        elif y_index==1:
            y_is = "Hold"
        elif y_index==2:
            y_is = "Sell"
        else:
            print("Some issue in counting y values")
    elif total_decisions==2:
        if y_index==0:
            y_is = "Buy"
        elif y_index==1:
            y_is = "Sell"
        else:
            print("Some issue in counting y values")
    else:
        print("Some issue in counting y values")

    i = 0
    counted_value = 0
    for each in y:
        if each[y_index] == 0:
            pass
        elif each[y_index] == 1:
            counted_value+=1
    print("Total {} after sampling: {}".format(y_is, counted_value))

def export_model(variable_list, pkl_filename, model_name):
    list_col = ['symbol','column_name','lookbacklen','period', 'movement',
                'accuracy','success_buy', 'success_sell', 'pkl_filename','model_used',
                'use_model','balanced','choice','frequency','timeframe','diff_type','docx_log_files_location']

    to_date = datetime.now()
    today_long = str(to_date.strftime("%Y-%m-%d %H %M %S"))

    log_model_df = pd.DataFrame(columns = list_col,index=[1])
    log_model_df['symbol'] = variable_list[0]
    log_model_df['column_name']=variable_list[1]
    log_model_df['lookbacklen']=variable_list[2]
    log_model_df['period']=variable_list[3]
    log_model_df['movement']=variable_list[4]
    log_model_df['success_buy']=int(variable_list[5]*100)
    log_model_df['success_sell']=int(variable_list[6]*100)
    log_model_df['accuracy']=variable_list[7]
    log_model_df['pkl_filename']=variable_list[8]
    log_model_df['model_used'] = str(variable_list[9])
    log_model_df['use_model'] = 'yes'
    log_model_df['balanced'] = variable_list[10]
    log_model_df['choice'] = variable_list[11]
    log_model_df['frequency'] = variable_list[12]
    log_model_df['timeframe'] = variable_list[13]
    log_model_df['diff_type'] = variable_list[14]
    log_model_df['docx_log_files_location'] = variable_list[15]


    with open(pkl_filename, 'wb') as file:
        pickle.dump(model_name, file)

    log_file_location = 'All_Exports/03_Model_Logs/model_log.csv'
    try:
        log_file_df = pd.read_csv(log_file_location)
        log_file_df = log_file_df.append(log_model_df)
        log_file_df.to_csv(log_file_location, index=False)
    except:
        log_model_df.to_csv(log_file_location, index=False)


def give_prediction(choose_model, input_file, log_file_df):
    choose_model = choose_model
    symbol = log_file_df['symbol'][choose_model]
    pkl_filename = log_file_df['pkl_filename'][choose_model]
    movement = log_file_df['movement'][choose_model]
    column_name = log_file_df['column_name'][choose_model]
    lookbacklen = log_file_df['lookbacklen'][choose_model]
    timeframe = log_file_df['timeframe'][choose_model]
    period = log_file_df['period'][choose_model]
    buy_acc = log_file_df['success_buy'][choose_model]
    sell_acc = log_file_df['success_sell'][choose_model]
    frequency = log_file_df['frequency'][choose_model]
    diff_type = log_file_df['diff_type'][choose_model]
    docx_log_files_location = log_file_df['docx_log_files_location'][choose_model]

    print('Model Location: {}'.format(pkl_filename))
    print('Model Log Docx Location: {}'.format(docx_log_files_location))

    how_long = 'in ' + str(round((period * 5) / 60, 1)) + ' hours'

    # Load from file
    with open(pkl_filename, 'rb') as file:
        pickle_model = pickle.load(file)

    filename = input_file
    load_data = pd.read_csv(filename)
    live_data = load_data.copy()
    live_data = live_data[['Date', column_name]].drop_duplicates()
    live_data = live_data.rename(columns={'Date': 'date'}).reset_index(drop=True)
    temp_copy = live_data.copy()
    entry_time = temp_copy['date'][temp_copy.shape[0] - 1]

    if diff_type == "yes":
        live_data[column_name] = live_data[column_name].pct_change(period)
        live_data = live_data.iloc[period:].reset_index(drop=True)
        wrt_date = temp_copy['date'][temp_copy.shape[0] - (1 + period)]
        wrt_price = temp_copy[column_name][temp_copy.shape[0] - (1+period)]
        calculated_level_buy = int(wrt_price*(1+movement))
        calculated_level_sell = int(wrt_price*(1 - movement))
    elif diff_type == "no":
        wrt_price = temp_copy[column_name][temp_copy.shape[0] - 1]
        calculated_level_buy = int(wrt_price + movement)
        calculated_level_sell = int(wrt_price - movement)
        pass
    else:
        print("Error processing input data")

    data_test = live_data.filter([column_name])
    dataset_test = data_test.values
    test_data = dataset_test[lookbacklen * -1:]

    x_predict_live = []
    x_predict_live.append(test_data)
    x_predict_live = np.array(x_predict_live)
    x_predict_live = np.reshape(x_predict_live, (1, lookbacklen))

    prediction = pickle_model.predict(x_predict_live)
    int(prediction)

    if int(prediction) == 0:
        action = 'BUY'
        print(
            "Symbol : {} ({})"
            "\nEntry Time: {}"
            "\nExit in: {} days"
            "\nAction: {}"
            "\nTouch level of {} expected in {} days"
            "\nWRT Price: {}"
            .format(symbol, column_name, entry_time, period, action, calculated_level_buy, period, int(wrt_price)))
    elif int(prediction) == 1:
        action = 'Do Nothing'
        print(
            "Symbol : {} ({})"
            "\nAction: {}"
            .format(symbol, column_name, action))
    elif int(prediction) == 2:
        action = 'SELL'
        print(
            "Symbol : {} ({})"
            "\nEntry Time: {}"
            "\nExit in: {} days"
            "\nAction: {}"
            "\nMovement of {} expected in {} days"
            "\nWRT Price: {}"
            .format(symbol, column_name, entry_time, period, action, calculated_level_sell, period,int(wrt_price)))
    else:
        print('No Valid Prediction')


def buy_sell_acc(cm):
    try:
        correct_buy = cm[0][0]
        incorrect_buy = cm[2][0]+cm[1][0]
        if correct_buy == 0 & incorrect_buy == 0:
            success_buy = 0
        else:
            success_buy = correct_buy / (correct_buy + incorrect_buy)

        correct_sell = cm[2][2]
        incorrect_sell = cm[0][2]+cm[1][2]
        if correct_sell == 0 & incorrect_sell == 0:
            success_sell = 0
        else:
            success_sell = correct_sell / (correct_sell + incorrect_sell)
        print('Buy Signal Success (ignoring hold): {}%'.format(int(success_buy * 100)))
        print('Sell Signal Success (ignoring hold): {}%'.format(int(success_sell * 100)))
    except:
        correct_buy = cm[0][0]
        incorrect_buy = cm[1][0]
        success_buy = correct_buy / (correct_buy + incorrect_buy)
        correct_sell = cm[1][1]
        incorrect_sell = cm[0][1]
        success_sell = correct_sell / (correct_sell + incorrect_sell)
        print('Buy Signal Success: {}%'.format(int(success_buy * 100)))
        print('Sell Signal Success: {}%'.format(int(success_sell * 100)))

    return success_buy, success_sell

def insert_df(document, df):
    t = document.add_table(df.shape[0] + 1, df.shape[1])

    # Setting Header
    for j in range(df.shape[-1]):
        t.cell(0, j).text = df.columns[j]

    # Setting data points
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i + 1, j).text = str(df.values[i, j])

def export_model_log(output_file_name, variable_list, cm_cr):
    cm_raw = cm_cr[0]
    cr_dict = cm_cr[1]

    cr = pd.DataFrame(cr_dict).transpose().reset_index()
    cr = cr.rename(columns={'index': 'CR'})
    cr = cr.round(2)

    cm = pd.DataFrame(cm_raw)
    cm = cm.rename(columns={0: 'Buy', 1: 'Hold', 2: 'Sell'})
    cm = cm.rename(index={0: 'Buy', 1: 'Hold', 2: 'Sell'})
    cm = cm.reset_index()
    cm = cm.rename(columns={'index': 'CM'})

    document = Document()
    section = document.sections[0]
    section.page_height = Mm(400)
    section.page_width = Mm(250)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)

    to_date = datetime.now()
    today_long = str(to_date.strftime("%Y-%m-%d"))

    run = document.add_paragraph().add_run()
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(5)
    font.italic = True

    # Heading
    document.add_heading('Model Vitals ', level=1)

    # Important information
    document.add_heading('Model information', level=2)
    document.add_paragraph('Name: ' + str(variable_list[9]))
    document.add_paragraph('Creation Date: ' + str(today_long) + ' (yyyy-mm-dd)')
    document.add_paragraph('Location: ' + str(variable_list[8]))

    document.add_heading('Other information', level=2)
    document.add_paragraph('Symbol: ' + str(variable_list[0]))
    document.add_paragraph('Series: ' + str(variable_list[1]))
    document.add_paragraph('Lookback: ' + str(variable_list[2]))
    document.add_paragraph('Period: ' + str(variable_list[3]))
    document.add_paragraph('Movement: ' + str(variable_list[4]))
    document.add_paragraph('Data dalancing done?: ' + str(variable_list[10]))
    document.add_paragraph('Sampling Type: ' + str(variable_list[11]))
    document.add_paragraph('Model on differences: ' + str(variable_list[14]))
    document.add_paragraph('Timeframe: ' + str(variable_list[13]))
    document.add_paragraph('Signal Giving Frequency: ' + str(variable_list[12]))

    # Model Evaluation
    document.add_heading('Model Performance', level=2)
    document.add_paragraph('Overall Accuracy: ' + str(int(100 * variable_list[5])) + '%')
    document.add_paragraph('Buy Accuracy: ' + str(int(variable_list[6] * 100)) + '%')
    document.add_paragraph('Sell Accuracy: ' + str(int(variable_list[7] * 100)) + '%')

    document.add_heading('Confusion Matrix', level=3)
    insert_df(document, cm)
    document.add_heading('Classification Report', level=3)
    insert_df(document, cr)

    # Saving the file
    # Save docx
    document.save(output_file_name)

    # Save HTML
    # with open(output_file_name, "rb") as docx_file:
    #     result = mammoth.convert_to_html(docx_file)
    #     html = result.value
    #
    # Html_file = open(output_file_name + '.html', "w")
    # Html_file.write(html)
    # Html_file.close()